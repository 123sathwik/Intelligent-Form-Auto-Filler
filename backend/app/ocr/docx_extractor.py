import logging
from pathlib import Path

logger = logging.getLogger("autofiller-backend")

class DOCXExtractor:
    @staticmethod
    def extract(file_path: Path) -> dict:
        """
        Extracts text from a DOCX file including headings, paragraphs, and tables.
        """
        import docx  # python-docx (lazy import)
        
        try:
            logger.info(f"Extracting DOCX file via python-docx: {file_path}")
            doc = docx.Document(str(file_path))
            
            combined_elements = []

            # 1. Extract paragraphs (includes headings since they are paragraphs with style)
            for para in doc.paragraphs:
                if para.text.strip():
                    combined_elements.append(para.text)

            # 2. Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            # Remove duplicate texts in cells caused by merged cells
                            val = cell.text.strip()
                            if not row_text or row_text[-1] != val:
                                row_text.append(val)
                    if row_text:
                        combined_elements.append(" | ".join(row_text))

            full_text = "\n\n".join(combined_elements)
            
            return {
                "text": full_text,
                "pages": 1,  # DOCX doesn't expose physical page numbers easily, default to 1
                "metadata": {},
                "extractor": "python-docx"
            }
        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}")
            raise e
